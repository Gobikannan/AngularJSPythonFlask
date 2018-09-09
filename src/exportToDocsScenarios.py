from docx import Document
from docx.shared import Inches
import datetime
import config

def save(id, scenarioName, releaseName, tests):
    
    appName = scenarioName + "_" + releaseName

    document = Document()
    dest_filename = config.FILE_CONFIG['folderPath'] + appName + "_" + id + ".docx"
    
    currentDT = datetime.datetime.now()
    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Scenario Name"
    hdr_cells[1].text = scenarioName
    row_cells = table.add_row().cells
    row_cells[0].text = "Release Name"
    row_cells[1].text = releaseName
    row_cells = table.add_row().cells
    row_cells[0].text = "Date"
    row_cells[1].text = currentDT.strftime("%d-%b-%Y")
    table.style = 'TableGrid'
    
    document.add_heading('Test Cases', level=2)

    table1 = document.add_table(rows=1, cols=5)
    table1.style = 'TableGrid'
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = "Type"
    hdr_cells[1].text = "Step"
    hdr_cells[2].text = "Data"
    hdr_cells[3].text = "Result"
    hdr_cells[4].text = "Comments"

    for test in tests:
        row_cells = table1.add_row().cells
        if(test['isheader'] == True):
            row_cells[0].text = test['filename']
        else:
            row_cells[0].text = test['type']
            row_cells[1].text = test['step']
            row_cells[2].text = test['data']
            row_cells[3].text = test['result']
            row_cells[4].text = test['comments']

    document.save(dest_filename)
