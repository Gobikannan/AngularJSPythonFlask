from docx import Document
from docx.shared import Inches
import datetime
import config

def save(id, target, freeText, ptasks, appName):

    document = Document()
    dest_filename = config.FILE_CONFIG['folderPath'] + appName + "_" + id + ".docx"
    
    currentDT = datetime.datetime.now()
    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Application Name"
    hdr_cells[1].text = appName
    row_cells = table.add_row().cells
    row_cells[0].text = "Target/Location"
    row_cells[1].text = target
    row_cells = table.add_row().cells
    row_cells[0].text = "Pre-Conditions, Environment Details"
    row_cells[1].text = freeText
    row_cells = table.add_row().cells
    row_cells[0].text = "Date"
    row_cells[1].text = currentDT.strftime("%d-%b-%Y")
    table.style = 'TableGrid'
    
    document.add_heading('Test Cases', level=2)

    table1 = document.add_table(rows=1, cols=4)
    table1.style = 'TableGrid'
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = "Sl #"
    hdr_cells[1].text = "Type"
    hdr_cells[2].text = "Step"
    hdr_cells[3].text = "Result"

    rowIndex = 1
    for ptask in ptasks:
        steps = ""
        for step in ptask['Steps']:
            steps = steps + step
            if(len(ptask['Steps']) > 1):
                steps = steps + ", "
        row_cells = table1.add_row().cells
        row_cells[0].text = str(rowIndex)
        row_cells[1].text = ptask['Type']
        row_cells[2].text = steps
        row_cells[3].text = ''
        rowIndex += 1

    document.save(dest_filename)
